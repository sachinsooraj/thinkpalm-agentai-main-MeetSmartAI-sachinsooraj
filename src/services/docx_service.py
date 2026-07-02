"""
MeetSmart AI — Word document (MoM) generator.
Uses python-docx to produce a branded ThinkPalm Minutes of Meeting document.
"""

from datetime import datetime
from pathlib import Path
from typing import List, Optional, Dict

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from src.utils.config import settings


# ThinkPalm brand colours
COLOR_PRIMARY = RGBColor(0x7C, 0x3A, 0xED)    # Purple
COLOR_DARK    = RGBColor(0x0F, 0x17, 0x2A)    # Dark navy
COLOR_ACCENT  = RGBColor(0xF5, 0x9E, 0x0B)    # Amber
COLOR_TEXT    = RGBColor(0x1E, 0x29, 0x3B)    # Slate
COLOR_LIGHT   = RGBColor(0x94, 0xA3, 0xB8)    # Light slate


def _set_cell_bg(cell, hex_color: str):
    """Set table cell background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_horizontal_line(doc: Document, color: str = "7C3AED"):
    """Add a coloured horizontal rule."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(0)
    return p


def build_mom_document(
    meeting_title: str,
    meeting_date: datetime,
    organizer_name: str,
    participants: List[Dict],   # [{"name": ..., "email": ..., "role": ...}]
    agenda: str,
    summary: str,
    decisions: List[str],
    action_items: List[Dict],   # [{"description", "owner", "deadline", "priority", "status"}]
    key_topics: Optional[List[str]] = None,
    save_path: Optional[Path] = None,
) -> Path:
    """
    Generate a formal MoM Word document with ThinkPalm branding.
    Returns the path to the saved .docx file.
    """
    doc = Document()

    # ── Page margins ──────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Header: Logo + Company name ───────────────────────────
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("🟣  THINKPALM TECHNOLOGIES")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = COLOR_PRIMARY

    sub_run = doc.add_paragraph()
    sub_run.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub_run.add_run("MeetSmart AI · Internal Meeting Management Platform")
    r.font.size = Pt(10)
    r.font.color.rgb = COLOR_LIGHT

    _add_horizontal_line(doc)

    # ── Document title ────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(12)
    t_run = title_p.add_run("MINUTES OF MEETING")
    t_run.bold = True
    t_run.font.size = Pt(22)
    t_run.font.color.rgb = COLOR_PRIMARY

    meeting_title_p = doc.add_paragraph()
    meeting_title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mt_run = meeting_title_p.add_run(meeting_title)
    mt_run.bold = True
    mt_run.font.size = Pt(15)
    mt_run.font.color.rgb = COLOR_TEXT

    _add_horizontal_line(doc)

    # ── Meeting metadata table ────────────────────────────────
    doc.add_paragraph()
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.style = "Table Grid"
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    meta_data = [
        ("Date & Time", meeting_date.strftime("%A, %d %B %Y  |  %I:%M %p IST")),
        ("Organizer",   organizer_name),
        ("Location",    "Conference Room / Google Meet"),
        ("Document ID", f"MOM-{meeting_date.strftime('%Y%m%d')}-{meeting_title[:20].replace(' ','_').upper()}"),
    ]
    for i, (label, value) in enumerate(meta_data):
        label_cell = meta_table.rows[i].cells[0]
        value_cell = meta_table.rows[i].cells[1]
        _set_cell_bg(label_cell, "1E293B")
        label_cell.width = Inches(1.8)
        label_p = label_cell.paragraphs[0]
        label_run = label_p.add_run(label)
        label_run.bold = True
        label_run.font.color.rgb = RGBColor(0xF1, 0xF5, 0xF9)
        label_run.font.size = Pt(10)
        value_p = value_cell.paragraphs[0]
        value_p.add_run(value).font.size = Pt(10)

    doc.add_paragraph()

    # ── Section: Participants ─────────────────────────────────
    _section_heading(doc, "1. PARTICIPANTS")
    ptable = doc.add_table(rows=1, cols=3)
    ptable.style = "Table Grid"
    for j, header_text in enumerate(["Name", "Role", "Email"]):
        cell = ptable.rows[0].cells[j]
        _set_cell_bg(cell, "7C3AED")
        run = cell.paragraphs[0].add_run(header_text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)

    for p in participants:
        row = ptable.add_row()
        row.cells[0].paragraphs[0].add_run(p.get("name", "")).font.size = Pt(10)
        row.cells[1].paragraphs[0].add_run(p.get("role", "")).font.size = Pt(10)
        row.cells[2].paragraphs[0].add_run(p.get("email", "")).font.size = Pt(10)

    doc.add_paragraph()

    # ── Section: Agenda ───────────────────────────────────────
    _section_heading(doc, "2. AGENDA")
    agenda_p = doc.add_paragraph()
    for line in agenda.strip().split("\n"):
        if line.strip():
            agenda_p = doc.add_paragraph(style="List Bullet")
            agenda_p.add_run(line.strip()).font.size = Pt(10)

    doc.add_paragraph()

    # ── Section: Summary ─────────────────────────────────────
    _section_heading(doc, "3. MEETING SUMMARY")
    summary_p = doc.add_paragraph()
    summary_p.add_run(summary).font.size = Pt(10)
    doc.add_paragraph()

    # ── Section: Key Decisions ───────────────────────────────
    _section_heading(doc, "4. KEY DECISIONS")
    if decisions:
        for i, decision in enumerate(decisions, 1):
            dp = doc.add_paragraph(style="List Number")
            dp.add_run(decision).font.size = Pt(10)
    else:
        doc.add_paragraph("No formal decisions recorded.").runs[0].font.size = Pt(10)
    doc.add_paragraph()

    # ── Section: Key Topics ──────────────────────────────────
    if key_topics:
        _section_heading(doc, "5. KEY TOPICS DISCUSSED")
        topics_para = doc.add_paragraph()
        topics_para.add_run(", ".join(key_topics)).font.size = Pt(10)
        doc.add_paragraph()

    # ── Section: Action Items ─────────────────────────────────
    section_num = 6 if key_topics else 5
    _section_heading(doc, f"{section_num}. ACTION ITEMS")
    if action_items:
        atable = doc.add_table(rows=1, cols=5)
        atable.style = "Table Grid"
        for j, header_text in enumerate(["#", "Action Item", "Owner", "Deadline", "Priority"]):
            cell = atable.rows[0].cells[j]
            _set_cell_bg(cell, "F59E0B")
            run = cell.paragraphs[0].add_run(header_text)
            run.bold = True
            run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
            run.font.size = Pt(10)

        for idx, item in enumerate(action_items, 1):
            row = atable.add_row()
            row.cells[0].paragraphs[0].add_run(str(idx)).font.size = Pt(10)
            row.cells[1].paragraphs[0].add_run(item.get("description", "")).font.size = Pt(10)
            row.cells[2].paragraphs[0].add_run(item.get("owner", "TBD")).font.size = Pt(10)
            row.cells[3].paragraphs[0].add_run(item.get("deadline", "TBD")).font.size = Pt(10)
            p_run = row.cells[4].paragraphs[0].add_run(item.get("priority", "medium").upper())
            p_run.font.size = Pt(10)
            priority_colors = {"HIGH": "EF4444", "MEDIUM": "F59E0B", "LOW": "22C55E"}
            p_color = priority_colors.get(item.get("priority", "medium").upper(), "F59E0B")
            _set_cell_bg(row.cells[4], p_color)
    else:
        doc.add_paragraph("No action items recorded.").runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # ── Footer ────────────────────────────────────────────────
    _add_horizontal_line(doc)
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run(
        f"Document generated by MeetSmart AI · ThinkPalm Technologies · "
        f"{datetime.utcnow().strftime('%d %B %Y, %H:%M UTC')}"
    )
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = COLOR_LIGHT
    footer_run.italic = True

    # ── Save ──────────────────────────────────────────────────
    if save_path is None:
        outputs_dir = settings.outputs_path
        safe_title = "".join(c if c.isalnum() or c in "_ " else "_" for c in meeting_title)
        filename = f"MoM_{meeting_date.strftime('%Y-%m-%d')}_{safe_title[:40]}.docx"
        save_path = outputs_dir / filename

    save_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(save_path))
    return save_path


def _section_heading(doc: Document, text: str):
    """Add a styled section heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = COLOR_PRIMARY


def build_sample_mom() -> Path:
    """Generate a sample MoM document for demo purposes."""
    from datetime import timedelta
    meeting_date = datetime(2026, 6, 20, 10, 0, 0)
    return build_mom_document(
        meeting_title="Q3 Product Roadmap Review",
        meeting_date=meeting_date,
        organizer_name="Priya Nair",
        participants=[
            {"name": "Priya Nair",     "role": "Product Manager",          "email": "priya.nair@thinkpalm.com"},
            {"name": "Arjun Sharma",   "role": "Senior Software Engineer", "email": "arjun.sharma@thinkpalm.com"},
            {"name": "Sanjay Pillai",  "role": "Engineering Manager",      "email": "sanjay.pillai@thinkpalm.com"},
            {"name": "Divya Krishnan", "role": "UX Designer",              "email": "divya.krishnan@thinkpalm.com"},
        ],
        agenda="1. Review Q2 deliverables\n2. Prioritise Q3 features\n3. Resource planning and allocation\n4. Timeline and milestone agreement",
        summary=(
            "The team reviewed Q2 performance metrics showing 85% milestone completion. "
            "Q3 priorities were discussed and aligned around three key themes: AI feature rollout, "
            "platform scalability improvements, and UX redesign. Resource allocation was finalised "
            "with Arjun leading the backend work and Divya leading the design sprint."
        ),
        decisions=[
            "AI feature development will begin in Week 1 of Q3 with Arjun's team.",
            "Design sprint for UX redesign scheduled for June 28 – July 5.",
            "Rahul will set up the staging environment by June 25.",
            "Weekly sync meetings every Tuesday at 10 AM IST confirmed.",
        ],
        action_items=[
            {"description": "Set up Q3 project board in Jira",       "owner": "Arjun Sharma",   "deadline": "2026-06-22", "priority": "high"},
            {"description": "Complete UX wireframes for new dashboard","owner": "Divya Krishnan","deadline": "2026-07-05", "priority": "high"},
            {"description": "Configure staging server environment",   "owner": "Rahul Menon",    "deadline": "2026-06-25", "priority": "medium"},
            {"description": "Draft Q3 OKRs document",                 "owner": "Priya Nair",    "deadline": "2026-06-24", "priority": "high"},
            {"description": "Share Q2 retrospective notes with team", "owner": "Sanjay Pillai", "deadline": "2026-06-21", "priority": "low"},
        ],
        key_topics=["Q3 Roadmap", "AI Features", "UX Redesign", "Resource Planning", "Staging Environment"],
        save_path=Path("samples/sample_mom.docx"),
    )
